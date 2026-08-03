"""
Document parsing service for PDF and DOCX files
"""

import io
import hashlib
from pathlib import Path
from typing import Tuple, Optional
import asyncio
from concurrent.futures import ThreadPoolExecutor

import PyPDF2
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
from loguru import logger

from ..config import settings
from ..core.exceptions import ParsingError, ValidationError
from ..models.enums import DocumentType

class DocumentParser:
    """Parse PDF and DOCX documents to extract text"""
    
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=2)
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024
        
    async def parse(self, file_path: Path, original_filename: str) -> Tuple[str, DocumentType, str]:
        """
        Parse document and extract text
        
        Returns:
            Tuple of (text, document_type, file_hash)
        """
        # Validate file
        if not file_path.exists():
            raise ParsingError(f"File not found: {file_path}")
        
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            raise ValidationError(f"File size exceeds {settings.max_file_size_mb}MB limit")
        
        # Calculate file hash
        file_hash = await self._compute_hash(file_path)
        
        # Determine document type
        doc_type = self._get_document_type(original_filename)
        
        # Extract text based on type
        try:
            if doc_type == DocumentType.PDF:
                text = await self._parse_pdf(file_path)
            elif doc_type == DocumentType.DOCX:
                text = await self._parse_docx(file_path)
            else:
                raise ParsingError(f"Unsupported document type: {original_filename}")
            
            # Validate extracted text
            if not text or len(text.strip()) < 50:
                logger.warning(f"Extracted very little text from {original_filename}: {len(text)} chars")
                if settings.ocr_enabled and doc_type == DocumentType.PDF:
                    logger.info(f"Attempting OCR fallback for {original_filename}")
                    text = await self._parse_pdf_with_ocr(file_path)
            
            return text.strip(), doc_type, file_hash
            
        except Exception as e:
            logger.error(f"Error parsing {original_filename}: {str(e)}")
            raise ParsingError(f"Failed to parse document: {str(e)}")
    
    async def _compute_hash(self, file_path: Path) -> str:
        """Compute SHA-256 hash of file"""
        def compute():
            sha256 = hashlib.sha256()
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(8192), b''):
                    sha256.update(chunk)
            return sha256.hexdigest()
        
        return await asyncio.get_event_loop().run_in_executor(
            self.executor, compute
        )
    
    async def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF using pdfplumber (better than PyPDF2)"""
        def parse():
            text_parts = []
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                        
                        # Also extract tables if present
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                if row:
                                    text_parts.append(' '.join(str(cell) for cell in row if cell))
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}, falling back to PyPDF2")
                # Fallback to PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text_parts.append(page.extract_text())
            
            return '\n'.join(text_parts)
        
        return await asyncio.get_event_loop().run_in_executor(
            self.executor, parse
        )
    
    async def _parse_pdf_with_ocr(self, file_path: Path) -> str:
        """Parse PDF using OCR for scanned documents"""
        if not settings.ocr_enabled:
            return ""
        
        def ocr_parse():
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text_parts = []
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    text = pytesseract.image_to_string(img)
                    text_parts.append(text)
                
                return '\n'.join(text_parts)
            except ImportError:
                logger.error("PyMuPDF not installed for OCR fallback")
                return ""
            except Exception as e:
                logger.error(f"OCR parsing failed: {e}")
                return ""
        
        return await asyncio.get_event_loop().run_in_executor(
            self.executor, ocr_parse
        )
    
    async def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX document"""
        def parse():
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
        
        return await asyncio.get_event_loop().run_in_executor(
            self.executor, parse
        )
    
    def _get_document_type(self, filename: str) -> DocumentType:
        """Determine document type from extension"""
        ext = Path(filename).suffix.lower()
        if ext == '.pdf':
            return DocumentType.PDF
        elif ext == '.docx':
            return DocumentType.DOCX
        else:
            return DocumentType.UNKNOWN